# -*- coding: utf-8 -*-
"""
Created on Thu Aug 19 10:52:22 2021

@author: VR787FC
"""

from docx2python import docx2python
import re
import xlwings
from docx import Document
myfile = Document('C:/Users/VR787FC/Downloads/20 Chagall Holdings Limited - AFS 2020 (EY 16 Aug 2021).docx')
for linenum,c in enumerate(myfile.paragraphs):
    print(linenum,c.text)
for i in myfile.paragraphs[984].paragraph_format.tab_stops:
    print((i.position, i.alignment))

print(myfile.paragraphs[984].text)


# myfile = docx2python('C:/Users/VR787FC/Downloads/20 Chagall Holdings Limited - AFS 2020 (EY 16 Aug 2021).docx')
# xlapp = xlwings.App(visible=True)
# wb = xlapp.books[0]
# wb.sheets[0].range('A1').select()
# ib = 0

# for i,line in enumerate(myfile.body[0][0][0]):
#     line = line.replace(' ','')
#         # case 1: Name&year|Notes|2020|2019
#         # case 2: Name&year|2020
#         # case 3: 2020| 2019
#         # case 4: 2020
#     if re.search('[A-Za-z ]{1,}[2019]{0,4}[\t]{1,}.*[()0-9,\-]{1,}[\t]{1,}[()0-9,\-\.]{1,}',line):
#         getdata = re.findall('[A-Za-z ]{1,}[2019]{0,4}[\t]{1,}.*[()0-9,\-]{1,}[\t]{1,}[()0-9,\-\.]{1,}',line.replace(' ',''))[0]
#         print(getdata.split('\t'))
#         if ib >= 10:
#             newsheet = wb.sheets.add()
#             newsheet.range('A1').select()
#         ib = 0
#         wb.app.selection.value = getdata.split('\t')
#         wb.app.selection.offset(1,0).select()
#     elif re.search('[A-Za-z ]{1,}[2019]{0,4}[\t]+[()0-9,\-\.]{1,}$',line):
#         getdata = re.findall('[A-Za-z ]{1,}[2019]{0,4}[\t]+[()0-9,\-\.]{1,}$',line)[0]
#         print(getdata.split('\t'))
#         if ib >= 10:
#             newsheet = wb.sheets.add()
#             newsheet.range('A1').select()
#         ib = 0
#         wb.app.selection.value = getdata.split('\t')
#         wb.app.selection.offset(1,0).select()
#     elif re.search('[()0-9,\-\.]{1,}[\t]+[()0-9,\-\.]{1,}$',line):
#         getdata = re.findall('[()0-9,\-\.]{1,}[\t]+[()0-9,\-\.]{1,}$',line)[0]
#         print(getdata.split('\t'))
#         if ib >= 10:
#             newsheet = wb.sheets.add()
#             newsheet.range('A1').select()
#         ib = 0
#         wb.app.selection.value = getdata.split('\t')
#         wb.app.selection.offset(1,0).select()
#     elif re.search('[\t]+([()0-9,\-\.]{1,}$)',line):
#         getdata = re.findall('[\t]+([()0-9,\-\.]{1,}$)',line)[0]
#         print(getdata.split('\t'))
#         if ib >= 10:
#             newsheet = wb.sheets.add()
#             newsheet.range('A1').select()
#         ib = 0
#         wb.app.selection.value = getdata.split('\t')
#         wb.app.selection.offset(1,0).select()
#     else:
#         ib = ib+1